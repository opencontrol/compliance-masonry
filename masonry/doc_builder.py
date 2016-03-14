""" This script uses uses core masonry objects to analyze missing parts of the
certification """
import os
from docx import Document

from masonry.core import Certification, Standard, Control

def document_reference(reference, document, certification_directory):
    if reference.get('type') == "Image":
        document.add_picture(
            os.path.join(certification_directory, reference.get('path'))
        )
    else:
        document.add_paragraph("{0} - {1}".format(
            reference.get('name'), reference.get('path')
        ))

class DocumentControl(Control):
    def prepare_narrative(self, narrative, document):
        if isinstance(narrative, dict):
            for key, text in narrative.items():
                document.add_paragraph("{0} - {1}".format(key, text))
        else:
            document.add_paragraph(narrative)

    def add_to_doc(self, document, get_verification, certification_directory):
        document.add_heading(self.meta.get('name'), 3)
        for component in self.justifications:
            system_key = component.get('system', 'No System')
            component_key = component.get('component', 'No Name')
            document.add_heading("{0} - {1}".format(system_key, component_key), 4)
            self.prepare_narrative(component.get('narrative'), document)
            for reference in component.get('references', []):
                document_reference(
                    get_verification(reference.get('system'), reference.get('component'), reference.get('verification')),
                    document,
                    certification_directory
                )

class DocumentStandard(Standard):
    def __init__(self, standards_yaml_path=None, standard_dict=None):
        super().__init__(
            standards_yaml_path=standards_yaml_path,
            standard_dict=standard_dict,
            control_class=DocumentControl
        )

    def add_to_doc(self, document, get_verification, certification_directory):
        for control_key, control in self:
            document.add_heading(str(control_key), 2)
            control.add_to_doc(document, get_verification, certification_directory)


class DocumentBuilder(Certification):
    """ InventoryBuilder load certification data and exports a yaml gap analysis """
    def __init__(self, certification_yaml_path):
        super().__init__(certification_yaml_path, standard_class=DocumentStandard)
        self.certification_directory = os.path.split(certification_yaml_path)[0]

    def document_component(self, component, document):
        document.add_heading(component.meta.get('name'), 2)
        for reference in component.meta.get('references', []):
            document_reference(reference, document, self.certification_directory)

    def document_systems(self, document):
        document.add_heading("Components", 0)
        for system_key, system in self.systems.items():
            document.add_heading(system.meta.get('name'), 1)
            for component in system:
                self.document_component(component, document)

    def document_standards(self, document):
        document.add_heading("Standards", 0)
        for standard_key, standard in self.standards_dict.items():
            document.add_heading(standard_key, 1)
            standard.add_to_doc(document, self.get_verification, self.certification_directory)

    def export(self, export_dir):
        export_path = os.path.join(export_dir, self.name + '.docx')
        document = Document()
        self.document_components(document)
        self.document_standards(document)
        document.save(export_path)
        return export_path
